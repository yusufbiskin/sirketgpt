from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO
import requests

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/olustur', methods=['POST'])
def olustur():
    ortak_sayisi = int(request.form.get('ortak_sayisi'))
    ortaklar = []
    toplam_sermaye = 0
    toplam_pay = 0

    for i in range(ortak_sayisi):
        kisi_tipi = request.form.get(f'ortak_kisi_tipi_{i}')
        ad = request.form.get(f'ortak_ad_{i}')
        tc = request.form.get(f'ortak_tc_{i}', '') if kisi_tipi == 'gercek' else ''
        vergi = request.form.get(f'ortak_vergi_{i}', '') if kisi_tipi == 'tuzel' else ''
        sermaye = int(request.form.get(f'ortak_sermaye_{i}') or 0)
        pay_adedi = int(request.form.get(f'ortak_pay_{i}') or 1)
        uyruk = request.form.get(f'ortak_uyruk_{i}')
        mudur = 'evet' if request.form.get(f'ortak_mudur_{i}') else 'hayir'
        imza = request.form.get(f'ortak_imza_{i}') or 'belirtilmedi'
        pay_tutari = round(sermaye / pay_adedi)

        ortaklar.append({
            'kisi_tipi': kisi_tipi,
            'ad': ad,
            'tc': tc,
            'vergi': vergi,
            'sermaye': sermaye,
            'pay_adedi': pay_adedi,
            'pay_tutari': pay_tutari,
            'uyruk': uyruk,
            'mudur': mudur,
            'imza': imza
        })

        toplam_sermaye += sermaye
        toplam_pay += pay_adedi

    unvan = request.form.get('unvan')
    merkez_il = request.form.get('merkez_il')
    merkez_ilce = request.form.get('merkez_ilce')
    adres = request.form.get('adres')
    faaliyet = request.form.get('faaliyet')

    headers = {
        "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY')}",
        "Content-Type": "application/json"
    }

    payload = {
    "model": "openai/gpt-3.5-turbo",
    "messages": [
        {
            "role": "system",
            "content": "Sen bir limited şirket ana sözleşmesi hazırlayan uzmansın. Yazacağın Madde 3 (Amaç ve Konu) metni MERSİS formatına %100 uygun olmalıdır. Resmi dil kullan ve numaralandırılmış şekilde alt başlıklar oluştur."
        },
        {
            "role": "user",
            "content": f"""{faaliyet} sektöründe faaliyet gösterecek bir limited şirket için aşağıdaki örneğe benzer şekilde Türk Ticaret Kanununa uygun Madde 3 - Amaç ve Konu metni hazırla:

	1. [Faaliyet Başlığı]
	1.1. ...
	1.2. ...
	2. [Faaliyet Başlığı]
	2.1. ...
	2.2. ...
	...
	Her faaliyet grubu için en az 5 alt madde olacak şekilde devam et. Maddeleri uzun cümleler kurarak resmi yazı dilinde ve kanunlara, mevzuata uygun şekilde yaz."""
        }
    ]
}



    try:
        response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
        amac_konu = response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        amac_konu = f"[GPT hatası]: {str(e)}"

    doc = Document()
    doc.add_heading('LİMİTED ŞİRKET ANA SÖZLEŞMESİ', 0)
    doc.add_heading('1. KURULUŞ', level=1)
    doc.add_paragraph("Aşağıdaki kurucular tarafından limited şirket kurulmuştur.")
    for o in ortaklar:
        kimlik = o['tc'] if o['kisi_tipi'] == 'gercek' else o['vergi']
        doc.add_paragraph(f"{o['ad']} - {kimlik} - {o['uyruk']}")

    doc.add_heading('2. UNVAN', level=1)
    doc.add_paragraph(f'Şirketin unvanı "{unvan} "dir.')

    doc.add_heading('3. AMAÇ VE KONU', level=1)
    doc.add_paragraph(amac_konu)

    doc.add_heading('4. ŞİRKETİN MERKEZİ', level=1)
    doc.add_paragraph(f"Sirketin merkezi {merkez_il.upper()} ili {merkez_ilce.upper()} ilçesi'dir.")
    doc.add_paragraph(f"Adresi {adres.upper()} 'dir.")
    doc.add_paragraph(
        "Adres degisikliginde yeni adres, ticaret siciline tescil ve Türkiye Ticaret Sicili Gazetesi'nde ilan ettirilir. "
        "Tescil ve ilan edilmis adrese yapılan tebligat sirkete yapılmıs sayılır. "
        "Tescil ve ilan edilmis adresinden ayrılmıs olmasına ragmen, yeni adresini süresi içinde tescil ettirmemis sirket için bu durum fesih sebebi sayılır."
    )



    doc.add_heading('5. SÜRE', level=1)
    doc.add_paragraph("Sirketin süresi, kurulusundan itibaren sınırsız'dır. Bu süre sirket sözlesmesini degistirmek suretiyle uzatılıp kısaltılabilir.")

    doc.add_heading('6. SERMAYE VE PAY SENETLERİNİN NEV’İ', level=1)
doc.add_paragraph(f"Sirketin sermayesi, beheri 25,00 Türk Lirası degerinde {toplam_pay} paya ayrılmıs toplam {toplam_sermaye:,.2f} Türk Lirası degerindedir.")

for o in ortaklar:
    doc.add_paragraph(
        f"-Beheri {o['pay_tutari']:,.2f} Türk Lirası değerinde {o['pay_adedi']} adet paya karşılık gelen "
        f"{o['sermaye']:,.2f} Türk Lirası {o['ad'].upper()} tarafından nakdi olarak taahhüt edilmiştir."
    )

doc.add_paragraph("Nakden taahhüt edilen payların itibari degerleri, sirketin tescilini izleyen 24 ay içinde ödenecektir.")


    doc.add_heading('7. ŞİRKETİN İDARESİ', level=1)
doc.add_paragraph("Sirketin isleri ve islemleri genel kurul tarafından seçilecek bir veya birkaç müdür tarafından yürütülür.")

for o in ortaklar:
    if o['mudur'] == 'evet':
        kimlik = o['tc'] if o['kisi_tipi'] == 'gercek' else o['vergi']
        doc.add_paragraph(
            f"Aksi Karar Alınıncaya Kadar {o['uyruk']} Uyruklu {kimlik} Kimlik No'lu , "
            f"{merkez_il.upper()} / {merkez_ilce.upper()} adresinde ikamet eden, {o['ad'].upper()} Müdür olarak seçilmistir."
        )
        doc.add_paragraph(f"Yetki Sekli: {o['imza'].capitalize()} Temsile Yetkilidir.")





    doc.add_heading('8. TEMSİL', level=1)
    doc.add_paragraph(
    "Sirketi müdürler temsil ederler. Sirketi temsil edecek imzalar genel kurul tarafından tespit, tescil ve ilan olunur. "
    "Müdürler, sirkete hizmet akdi ile baglı olanları sınırlı yetkiye sahip ticari vekil veya diger tacir yardımcıları olarak atayabilir. "
    "Bu sekilde atanacak olanların görev ve yetkileri, hazırlanacak iç yönergede açıkça belirlenir. Bu durumda iç yönergenin tescil ve ilanı zorunludur. "
    "Iç yönerge ile ticari vekil ve diger tacir yardımcıları atanamaz. Yetkilendirilen kişiler ticaret siciline tescil ve ilan edilir. "
    "Bu kisilerin, sirkete ve üçüncü kisilere verecekleri her tür zarardan dolayı müdürler müteselsilen sorumludur."
)


    doc.add_heading('9. GENEL KURUL', level=1)
    doc.add_paragraph(
    "Genel Kurullar, olagan ve olaganüstü toplanırlar. Olagan genel kurul, her yıl hesap döneminin sona ermesinden itibaren 3 ay içinde; "
    "olaganüstü genel kurullar ise, Sirket islerinin gerektirdigi hallerde ve zamanlarda toplanır. "
    "Genel kurul toplantılarında, her ortagın oy hakkı, esas sermaye paylarının itibari degerine göre hesaplanır. "
    "Genel kurul toplantıları ve bu toplantılardaki karar nisabı, Türk Ticaret Kanunu hükümlerine tabidir. "
    "Genel kurul, sirketin merkez adresinde veya yönetim merkezinin bulundugu sehrin elverisli bir yerinde toplanır."
)


    doc.add_heading('10. İLAN', level=1)
    doc.add_paragraph(
    "Genel kurulun toplantıya çagrılmasına iliskin ilanlar da dahil olmak üzere Sirkete ait ilanlar Türkiye Ticaret Sicili Gazetesinde yapılır. "
    "Genel kurul toplantılarına iliskin ilanların toplantı gününden en az on gün önce yapılması zorunludur."
)


    doc.add_heading('11. HESAP DÖNEMİ', level=1)
    doc.add_paragraph(
    "Sirketin hesap yılı, Ocak ayının 1. gününden baslar ve Aralık ayının 31. günü sona erer. "
    "Fakat birinci hesap yılı, Sirketin kesin olarak kuruldugu tarihten itibaren baslar ve o senenin aralık ayının otuz birinci günü sona erer."
)


    doc.add_paragraph(
    "Sirketin net dönem karı yapılmıs her çesit masrafların çıkarılmasından sonra kalan miktardır. "
    "Net dönem kârından her yıl %5 genel kanuni yedek akçe ayrılır; kalan miktar, genel kurul kararı ile pay sahiplerine kar payı olarak dagıtılır. "
    "Kar payı, esas sermaye payının itibari degerine, yerine getirilen ek ödeme yükümlülügünün tutarı eklenmek suretiyle olusacak toplam miktara oranla hesaplanır."
)


    doc.add_heading('13. YEDEK AKÇE', level=1)
    doc.add_paragraph("Yedek akçeler, Türk Ticaret Kanunu’nun 519-523. maddelerine göre ayrılır.")

    doc.add_heading('14. KANUNİ HÜKÜMLER', level=1)
    doc.add_paragraph("Bu sözleşmede hüküm bulunmayan hallerde Türk Ticaret Kanunu hükümleri uygulanır.")

    doc.add_paragraph("\nKurucular")
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Kurucu"
    table.rows[0].cells[1].text = "Uyruk"
    table.rows[0].cells[2].text = "İmza"
    for o in ortaklar:
        row = table.add_row().cells
        row[0].text = o['ad']
        row[1].text = o['uyruk']
        row[2].text = "________"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="sirket_ana_sozlesmesi.docx")

if __name__ == '__main__':
    app.run(debug=True)